#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Transkriptionstool für MP4-Dateien - Konsolenversion
Falls GUI nicht verfügbar ist, kann dieses Skript direkt verwendet werden.
"""

import os
import sys
import tempfile
from moviepy.editor import VideoFileClip
import speech_recognition as sr
from docx import Document
from pydub import AudioSegment


def extract_audio(mp4_path):
    """Extrahiert Audio aus MP4-Datei und speichert als WAV"""
    try:
        # Temporäre Datei erstellen
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_audio:
            temp_audio_path = temp_audio.name
        
        # Audio mit moviepy extrahieren
        video = VideoFileClip(mp4_path)
        audio = video.audio
        
        # Als WAV speichern
        audio.write_audiofile(temp_audio_path, codec='pcm_s16le')
        audio.close()
        video.close()
        
        return temp_audio_path
        
    except Exception as e:
        print(f"Fehler beim Extrahieren des Audios: {e}")
        return None


def audio_to_text(audio_path):
    """Konvertiert Audio-Datei in Text mit Spracherkennung"""
    recognizer = sr.Recognizer()
    
    try:
        # Audio-Datei laden
        audio_file = AudioSegment.from_wav(audio_path)
        
        # In Chunks aufteilen (für bessere Erkennung)
        chunk_size = 30 * 1000  # 30 Sekunden
        chunks = [audio_file[i:i+chunk_size] for i in range(0, len(audio_file), chunk_size)]
        
        full_transcript = ""
        
        for i, chunk in enumerate(chunks):
            # Temporäre WAV-Datei für den Chunk erstellen
            with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_chunk:
                chunk.export(temp_chunk.name, format="wav")
                
                # Text erkennen
                with sr.AudioFile(temp_chunk.name) as source:
                    audio_data = recognizer.record(source)
                    
                    try:
                        # Versuche mit Google Speech Recognition
                        text = recognizer.recognize_google(audio_data, language="de-DE")
                        full_transcript += text + " "
                        print(f"Chunk {i+1}/{len(chunks)}: {text}")
                        
                    except sr.UnknownValueError:
                        print(f"Chunk {i+1}: Spracherkennung fehlgeschlagen (unbekannter Wert)")
                    except sr.RequestError as e:
                        print(f"Chunk {i+1}: API-Fehler: {e}")
                        # Versuche mit Sphinx (offline)
                        try:
                            text = recognizer.recognize_sphinx(audio_data, language="de-de")
                            full_transcript += text + " "
                            print(f"Chunk {i+1} (Sphinx): {text}")
                        except Exception as sphinx_error:
                            print(f"Chunk {i+1}: Sphinx-Fehler: {sphinx_error}")
                
                # Temporäre Datei löschen
                os.unlink(temp_chunk.name)
        
        return full_transcript.strip()
        
    except Exception as e:
        print(f"Fehler bei der Spracherkennung: {e}")
        return None
    finally:
        # Temporäre Audio-Datei löschen
        if os.path.exists(audio_path):
            os.unlink(audio_path)


def save_as_word(text, output_path=None):
    """Speichert den Text als Word-Datei"""
    try:
        # Dokument erstellen
        doc = Document()
        doc.add_heading('Transkription', level=1)
        doc.add_paragraph(text)
        
        # Speicherort festlegen
        if not output_path:
            input_path = input("Geben Sie den Pfad zur MP4-Datei ein: ")
            if not os.path.exists(input_path):
                print("Datei nicht gefunden!")
                return None
            
            output_dir = os.path.dirname(input_path)
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}_transkription.docx")
        
        # Datei speichern
        doc.save(output_path)
        
        return output_path
        
    except Exception as e:
        print(f"Fehler beim Speichern der Word-Datei: {e}")
        return None


def main():
    """Hauptfunktion für die Konsolenversion"""
    print("MP4 Transkriptionstool - Konsolenversion")
    print("=" * 50)
    
    # Dateipfad abfragen
    mp4_path = input("Geben Sie den Pfad zur MP4-Datei ein: ").strip()
    
    if not os.path.exists(mp4_path):
        print("Fehler: Datei nicht gefunden!")
        return
    
    if not mp4_path.lower().endswith('.mp4'):
        print("Fehler: Bitte geben Sie eine MP4-Datei an!")
        return
    
    print(f"\nVerarbeite Datei: {mp4_path}")
    
    # Audio extrahieren
    print("Extrahiere Audio aus Video...")
    audio_path = extract_audio(mp4_path)
    
    if not audio_path:
        print("Fehler: Konnte Audio nicht extrahieren!")
        return
    
    # Audio transkribieren
    print("Transkribiere Audio...")
    transcript = audio_to_text(audio_path)
    
    if not transcript:
        print("Fehler: Konnte Audio nicht transkribieren!")
        return
    
    # Als Word speichern
    print("Speichere Word-Datei...")
    output_dir = os.path.dirname(mp4_path)
    base_name = os.path.splitext(os.path.basename(mp4_path))[0]
    output_path = os.path.join(output_dir, f"{base_name}_transkription.docx")
    
    doc = Document()
    doc.add_heading('Transkription', level=1)
    doc.add_paragraph(transcript)
    doc.save(output_path)
    
    print(f"\nFertig! Transkription gespeichert unter: {output_path}")


if __name__ == "__main__":
    main()
