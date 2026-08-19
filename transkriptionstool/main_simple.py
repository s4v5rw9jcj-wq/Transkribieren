#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Einfaches Transkriptionstool für MP4-Dateien
Erstellt eine GUI zum Hochladen von MP4-Dateien und Transkribieren in Word-Dateien.
"""

import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import tempfile
import threading
from moviepy.editor import VideoFileClip
import speech_recognition as sr
from docx import Document
from pydub import AudioSegment


class TranskriptionsApp:
    def __init__(self, root):
        self.root = root
        self.root.title("MP4 Transkriptionstool")
        self.root.geometry("600x400")
        self.root.resizable(False, False)
        
        # Variablen
        self.file_path = None
        self.is_processing = False
        self.recognizer = sr.Recognizer()
        
        # GUI erstellen
        self.create_widgets()
    
    def create_widgets(self):
        """Erstellt alle GUI-Elemente"""
        
        # Hauptframe
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Titel
        title_label = ttk.Label(
            main_frame, 
            text="MP4 Audio Transkriptionstool",
            font=('Helvetica', 16, 'bold')
        )
        title_label.pack(pady=(0, 20))
        
        # Datei-Upload-Bereich
        upload_frame = ttk.LabelFrame(main_frame, text="Datei hochladen", padding="15")
        upload_frame.pack(fill=tk.X, pady=(0, 20))
        
        # Datei auswählen Button
        browse_btn = ttk.Button(
            upload_frame,
            text="MP4-Datei auswählen",
            command=self.browse_file
        )
        browse_btn.pack(pady=10)
        
        # Datei-Info
        self.file_info = ttk.Label(main_frame, text="Keine Datei ausgewählt")
        self.file_info.pack(pady=(0, 20))
        
        # Fortschrittsbalken
        self.progress = ttk.Progressbar(main_frame, orient=tk.HORIZONTAL, length=500, mode='determinate')
        self.progress.pack(pady=(0, 20))
        self.progress['value'] = 0
        
        # Status-Label
        self.status_label = ttk.Label(main_frame, text="Bereit")
        self.status_label.pack(pady=(0, 10))
        
        # Button-Frame
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=(10, 0))
        
        # Transkribieren-Button
        self.transcribe_btn = ttk.Button(
            button_frame,
            text="Transkribieren starten",
            command=self.start_transcription_thread,
            state=tk.DISABLED
        )
        self.transcribe_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Abbrechen-Button
        self.cancel_btn = ttk.Button(
            button_frame,
            text="Abbrechen",
            command=self.cancel_transcription,
            state=tk.DISABLED
        )
        self.cancel_btn.pack(side=tk.LEFT)
        
        # Stilen anpassen
        self.style_widgets()
    
    def style_widgets(self):
        """Stilt die Widgets für besseres Aussehen"""
        style = ttk.Style()
        
        # Button-Stile
        style.configure('TButton', padding=6, font=('Helvetica', 10))
        
        # Label-Stile
        style.configure('TLabel', font=('Helvetica', 10))
        style.configure('TLabelFrame.Label', font=('Helvetica', 10, 'bold'))
    
    def browse_file(self):
        """Öffnet einen Datei-Dialog zum Auswählen einer MP4-Datei"""
        file_path = filedialog.askopenfilename(
            title="MP4-Datei auswählen",
            filetypes=[("MP4-Dateien", "*.mp4"), ("Alle Dateien", "*.*")]
        )
        
        if file_path:
            self.file_path = file_path
            self.update_file_info()
            self.transcribe_btn.config(state=tk.NORMAL)
    
    def update_file_info(self):
        """Aktualisiert die Datei-Informationen"""
        if self.file_path:
            file_name = os.path.basename(self.file_path)
            file_size = os.path.getsize(self.file_path) / (1024 * 1024)  # in MB
            self.file_info.config(
                text=f"Ausgewählte Datei: {file_name} ({file_size:.2f} MB)"
            )
    
    def start_transcription_thread(self):
        """Startet die Transkription in einem separaten Thread"""
        if not self.file_path:
            messagebox.showerror("Fehler", "Bitte wählen Sie zuerst eine MP4-Datei aus.")
            return
        
        # UI für Verarbeitung aktualisieren
        self.is_processing = True
        self.transcribe_btn.config(state=tk.DISABLED)
        self.cancel_btn.config(state=tk.NORMAL)
        self.status_label.config(text="Verarbeite Datei...")
        self.progress['value'] = 0
        
        # Thread starten
        threading.Thread(
            target=self.transcribe_file,
            daemon=True
        ).start()
    
    def cancel_transcription(self):
        """Bricht die Transkription ab"""
        self.is_processing = False
        self.status_label.config(text="Abbruch angefordert...")
        self.cancel_btn.config(state=tk.DISABLED)
    
    def transcribe_file(self):
        """Hauptmethode zur Transkription der Datei"""
        try:
            # 1. Audio aus MP4 extrahieren
            self.status_label.config(text="Extrahiere Audio aus Video...")
            self.progress['value'] = 10
            self.root.update()
            
            audio_path = self.extract_audio()
            if not audio_path:
                raise Exception("Konnte Audio nicht extrahieren")
            
            # 2. Audio in Text transkribieren
            self.status_label.config(text="Transkribiere Audio...")
            self.progress['value'] = 40
            self.root.update()
            
            transcript = self.audio_to_text(audio_path)
            if not transcript:
                raise Exception("Konnte Audio nicht transkribieren")
            
            # 3. Transkript als Word-Datei speichern
            self.status_label.config(text="Speichere Word-Datei...")
            self.progress['value'] = 80
            self.root.update()
            
            word_path = self.save_as_word(transcript)
            
            # 4. Aufräumen
            self.progress['value'] = 100
            self.status_label.config(text=f"Fertig! Datei gespeichert: {word_path}")
            self.transcribe_btn.config(state=tk.NORMAL)
            self.cancel_btn.config(state=tk.DISABLED)
            self.is_processing = False
            
            # Erfolgmeldung
            messagebox.showinfo(
                "Erfolg",
                f"Transkription erfolgreich!\n\nWord-Datei gespeichert unter:\n{word_path}"
            )
            
        except Exception as e:
            self.status_label.config(text=f"Fehler: {str(e)}")
            self.is_processing = False
            self.transcribe_btn.config(state=tk.NORMAL)
            self.cancel_btn.config(state=tk.DISABLED)
            messagebox.showerror("Fehler", f"Ein Fehler ist aufgetreten:\n{str(e)}")
        
        finally:
            self.is_processing = False
    
    def extract_audio(self):
        """Extrahiert Audio aus MP4-Datei und speichert als WAV"""
        try:
            # Temporäre Datei erstellen
            with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_audio:
                temp_audio_path = temp_audio.name
            
            # Audio mit moviepy extrahieren
            video = VideoFileClip(self.file_path)
            audio = video.audio
            
            # Als WAV speichern
            audio.write_audiofile(temp_audio_path, codec='pcm_s16le')
            audio.close()
            video.close()
            
            return temp_audio_path
            
        except Exception as e:
            print(f"Fehler beim Extrahieren des Audios: {e}")
            return None
    
    def audio_to_text(self, audio_path):
        """Konvertiert Audio-Datei in Text mit Spracherkennung"""
        try:
            # Audio-Datei laden
            audio_file = AudioSegment.from_wav(audio_path)
            
            # In Chunks aufteilen (für bessere Erkennung)
            chunk_size = 30 * 1000  # 30 Sekunden
            chunks = [audio_file[i:i+chunk_size] for i in range(0, len(audio_file), chunk_size)]
            
            full_transcript = ""
            
            for i, chunk in enumerate(chunks):
                if not self.is_processing:
                    break
                
                # Temporäre WAV-Datei für den Chunk erstellen
                with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_chunk:
                    chunk.export(temp_chunk.name, format="wav")
                    
                    # Text erkennen
                    with sr.AudioFile(temp_chunk.name) as source:
                        audio_data = self.recognizer.record(source)
                        
                        try:
                            # Versuche mit Google Speech Recognition
                            text = self.recognizer.recognize_google(audio_data, language="de-DE")
                            full_transcript += text + " "
                            
                        except sr.UnknownValueError:
                            print(f"Chunk {i+1}: Spracherkennung fehlgeschlagen (unbekannter Wert)")
                        except sr.RequestError as e:
                            print(f"Chunk {i+1}: API-Fehler: {e}")
                            # Versuche mit Sphinx (offline)
                            try:
                                text = self.recognizer.recognize_sphinx(audio_data, language="de-de")
                                full_transcript += text + " "
                            except Exception as sphinx_error:
                                print(f"Chunk {i+1}: Sphinx-Fehler: {sphinx_error}")
                    
                    # Temporäre Datei löschen
                    os.unlink(temp_chunk.name)
                
                # Fortschritt aktualisieren
                progress = 40 + int((i + 1) / len(chunks) * 40)
                self.progress['value'] = progress
                self.status_label.config(text=f"Transkribiere... ({i+1}/{len(chunks)} Chunks)")
                self.root.update()
            
            return full_transcript.strip()
            
        except Exception as e:
            print(f"Fehler bei der Spracherkennung: {e}")
            return None
        finally:
            # Temporäre Audio-Datei löschen
            if os.path.exists(audio_path):
                os.unlink(audio_path)
    
    def save_as_word(self, text):
        """Speichert den Text als Word-Datei"""
        try:
            # Dokument erstellen
            doc = Document()
            doc.add_heading('Transkription', level=1)
            doc.add_paragraph(text)
            
            # Speicherort festlegen
            output_dir = os.path.dirname(self.file_path)
            base_name = os.path.splitext(os.path.basename(self.file_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}_transkription.docx")
            
            # Datei speichern
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            print(f"Fehler beim Speichern der Word-Datei: {e}")
            return None


def main():
    """Hauptfunktion zum Starten der Anwendung"""
    root = tk.Tk()
    app = TranskriptionsApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
